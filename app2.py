import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
import base64
import google.generativeai as genai  # Added for Gemini Pro
import os  # For environment variables
from dotenv import load_dotenv  # Load API Key from .env file
import qrcode  # For QR Code Generation
from PIL import Image  # For Displaying QR Code
import io  # For handling image data
from urllib.parse import urlparse  # For LinkedIn URL validation
import requests  # For fetching LinkedIn data (if using a public API)

# Load API Key from .env file
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

if not GOOGLE_API_KEY:
    st.error("Please set the GOOGLE_API_KEY environment variable.")
    st.stop()


genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-pro')


def summarize_experience(description):
    """Summarizes a job description using Gemini Pro."""
    prompt = f"""Summarize the following job description in 3 concise bullet points, highlighting accomplishments and quantifiable results.  Focus on keywords that would be relevant to recruiters:\n\n{description}"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error summarizing experience: {e}"


def generate_skills_suggestions(description):
    """Suggests relevant skills based on the job description using Gemini Pro."""
    prompt = f"""Extract a list of relevant skills from the following job description:\n\n{description}"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error generating skills: {e}"


def generate_cover_letter(name, job_description):
    """Generates a personalized cover letter based on the job description."""
    prompt = f"""Write a cover letter for {name} based on the following job description.  The cover letter should be concise and highlight relevant skills and experience.  Make it professional and engaging:\n\n{job_description}"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error generating cover letter: {e}"

def suggest_improvements(job_description, skills):
    """Suggests improvements based on job description and current skills."""
    prompt = f"""Based on the following job description and provided skills, suggest improvements to strengthen the profile, including relevant certifications, projects, or portfolios to add: \n\nJob Description: {job_description}\nSkills: {skills}"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error generating suggestions: {e}"


def extract_skills_qualifications(job_description):
    """Extracts key skills, qualifications, and responsibilities from the job description."""
    prompt = f"""Extract key skills, qualifications, and responsibilities from the following job description:\n\n{job_description}"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error extracting skills and qualifications: {e}"


def highlight_missing_skills(user_skills, extracted_skills):
    """Highlights missing skills in the user's resume based on the job description."""
    prompt = f"""You are a career expert comparing two lists of skills. Given the following list of skills the candidate posesses: {user_skills}, and the following list of skills required by a job description: {extracted_skills}, what skills are the candidate missing?
    Make your response a comma separated list of skills"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error highlighting missing skills: {e}"


def generate_qr_code(data):
    """Generates a QR code image."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(data)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    # Convert PIL Image to BytesIO for Streamlit display
    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='PNG')
    img_byte_arr = img_byte_arr.getvalue()
    return img_byte_arr  # Return byte array

def is_valid_linkedin_url(url):
    """Validates a LinkedIn profile URL."""
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc == "www.linkedin.com", result.path.startswith("/in/")])
    except:
        return False

def fetch_linkedin_data(linkedin_url):
    """
    Fetches LinkedIn data.  *Important:* Directly scraping LinkedIn is against their terms of service and can lead to your IP being blocked.  This function is a placeholder and should only be used with a proper API or with extreme caution and user consent.
    """
    st.warning("Fetching LinkedIn data directly is risky.  Use a LinkedIn API or proceed with caution.")
    try:
        # This is a placeholder.  Replace with your actual LinkedIn API call or scraping logic
        # Be VERY careful and respect LinkedIn's terms of service.
        # For demonstration purposes, let's just return a dummy dataset.
        return {
            "name": "John Doe (Example)",
            "skills": "Python, Data Analysis, Machine Learning",
            "experience": [
                {"title": "Data Scientist", "company": "Example Corp", "dates": "2020-Present", "description": "Built machine learning models."},
            ],
            "education": [
                {"degree": "Master's in CS", "school": "Example University", "dates": "2018-2020", "description": "Focus on AI."}
            ]
        }
    except Exception as e:
        st.error(f"Error fetching LinkedIn data: {e}")
        return None


def suggest_profile_improvements(linkedin_data):
    """Suggests LinkedIn profile improvements based on industry benchmarks."""
    prompt = f"""Based on the following LinkedIn profile data, suggest improvements to the profile to make it stronger based on industry best practices. Consider profile completeness, keyword usage, and showcasing accomplishments:\n\n{linkedin_data}"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error suggesting profile improvements: {e}"


def generate_interview_questions(resume_content):
    """Generates possible interview questions based on the resume content."""
    prompt = f"""Generate potential interview questions based on the following resume content. Cover a range of topics, including experience, skills, and projects:\n\n{resume_content}"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error generating interview questions: {e}"


def analyze_strengths_weaknesses(resume_content):
    """Analyzes strengths and weaknesses based on the resume content."""
    prompt = f"""Analyze the strengths and weaknesses of a candidate based on the following resume content. Provide constructive feedback:\n\n{resume_content}"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error analyzing strengths and weaknesses: {e}"


def generate_resume(name, email, phone, linkedin, github, skills, experience, education, projects):
    """Generates a resume as a docx document."""

    document = Document()

    # Header Information
    document.add_heading(name, level=1)
    header_info = f"{email} | {phone} | [LinkedIn]({linkedin}) | [GitHub]({github})"
    document.add_paragraph(header_info)

    # Skills Section
    document.add_heading("Skills", level=2)
    skills_list = skills.split(",")  # Split comma-separated skills
    for skill in skills_list:
        document.add_paragraph(skill.strip(), style='List Bullet')

    # Experience Section
    document.add_heading("Experience", level=2)
    for exp in experience:
        document.add_paragraph(exp['title'], style='Heading 3')
        document.add_paragraph(exp['company'] + " | " + exp['dates'])
        # Use summarized experience if available
        if 'description' in exp and exp['description']:
            summarized_description = summarize_experience(exp['description'])
            document.add_paragraph(summarized_description)
        else:
            document.add_paragraph(exp['description'])


    # Education Section
    document.add_heading("Education", level=2)
    for edu in education:
        document.add_paragraph(edu['degree'], style='Heading 3')
        document.add_paragraph(edu['school'] + " | " + edu['dates'])
        document.add_paragraph(edu['description'])

    # Projects Section
    document.add_heading("Projects", level=2)
    for proj in projects:
        document.add_paragraph(proj['title'], style='Heading 3')
        document.add_paragraph(proj['description'])

    return document


def main():
    st.title("Smart Resume Generator")

    # User Input Form
    with st.form("resume_form"):
        st.header("Personal Information")
        name = st.text_input("Full Name")
        email = st.text_input("Email Address")
        phone = st.text_input("Phone Number")
        linkedin = st.text_input("LinkedIn Profile URL")
        github = st.text_input("GitHub Profile URL")

        st.header("Skills")
        skills = st.text_area("Skills (comma-separated)")

        st.header("Experience")
        experience = []
        num_experiences = st.number_input("Number of Experience Entries", min_value=0, max_value=5, value=1, step=1)
        for i in range(num_experiences):
            st.subheader(f"Experience #{i+1}")
            title = st.text_input(f"Job Title #{i+1}")
            company = st.text_input(f"Company #{i+1}")
            dates = st.text_input(f"Dates (e.g., Jan 2020 - Present) #{i+1}")
            description = st.text_area(f"Description of Responsibilities #{i+1}")
            skills_suggestions = generate_skills_suggestions(description)  # Generate skills suggestions
            st.write(f"Suggested Skills for this Experience:\n{skills_suggestions}")
            experience.append({"title": title, "company": company, "dates": dates, "description": description})



        st.header("Education")
        education = []
        num_educations = st.number_input("Number of Education Entries", min_value=0, max_value=3, value=1, step=1)
        for i in range(num_educations):
            st.subheader(f"Education #{i+1}")
            degree = st.text_input(f"Degree #{i+1}")
            school = st.text_input(f"School #{i+1}")
            dates = st.text_input(f"Dates (e.g., Sep 2016 - May 2020) #{i+1}")
            description = st.text_area(f"Description of Studies/Achievements #{i+1}")
            education.append({"degree": degree, "school": school, "dates": dates, "description": description})

        st.header("Projects")
        projects = []
        num_projects = st.number_input("Number of Projects", min_value=0, max_value=5, value=1, step=1)
        for i in range(num_projects):
            st.subheader(f"Project #{i+1}")
            title = st.text_input(f"Project Title #{i+1}")
            description = st.text_area(f"Project Description #{i+1}")
            projects.append({"title": title, "description": description})

        st.header("Job Description Analysis")
        job_description = st.text_area("Paste Job Description Here", help="Paste the job description to analyze for missing skills.")

        submitted = st.form_submit_button("Generate Resume")

    if submitted:
        # Input Validation (Basic) - Add more robust validation as needed.
        if not name or not email or not skills:
            st.error("Please fill in all required fields (Name, Email, Skills).")
        else:
            # LinkedIn Integration
            if linkedin and is_valid_linkedin_url(linkedin):
                linkedin_data = fetch_linkedin_data(linkedin)  # Remember the risks!
                if linkedin_data:
                    st.success("LinkedIn data fetched (with caution!).")
                    # Auto-fill form fields (example)
                    # name = linkedin_data.get("name", name)  # Use LinkedIn data if available
                    profile_improvements = suggest_profile_improvements(linkedin_data)
                    st.header("LinkedIn Profile Improvement Suggestions")
                    st.write(profile_improvements)
            else:
                st.warning("Invalid LinkedIn URL or LinkedIn integration skipped.")


            # Generate the Resume Document
            document = generate_resume(name, email, phone, linkedin, github, skills, experience, education, projects)

            # Convert the document to a string for analysis
            resume_text = ""
            for paragraph in document.paragraphs:
                resume_text += paragraph.text + "\n"

            # Download the Resume
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)  # Go back to the beginning of the buffer

            # Create download link
            b64 = base64.b64encode(buffer.read()).decode()
            href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{name}_resume.docx">Download Resume (docx)</a>'
            st.markdown(href, unsafe_allow_html=True)

            # QR Code Generation
            resume_url = f"https://linkedin.com/in/{linkedin}"  # Or create an actual interactive resume URL
            qr_code_image = generate_qr_code(resume_url)
            st.image(qr_code_image, caption="Scan to view interactive resume (placeholder)", use_column_width=False)

            # Cover Letter Generation
            cover_letter = generate_cover_letter(name, job_description)
            st.header("Generated Cover Letter")
            st.write(cover_letter)

            # Improvement Suggestions
            improvements = suggest_improvements(job_description, skills)
            st.header("Suggested Improvements")
            st.write(improvements)

            # Skills Extraction and Highlighting
            extracted_skills = extract_skills_qualifications(job_description)
            st.header("Extracted Skills from Job Description")
            st.write(extracted_skills)

            missing_skills = highlight_missing_skills(skills, extracted_skills)
            st.header("Missing Skills (Based on Job Description)")
            st.write(missing_skills)

            # Interview Preparation
            interview_questions = generate_interview_questions(resume_text)
            st.header("Possible Interview Questions")
            st.write(interview_questions)

            strengths_weaknesses = analyze_strengths_weaknesses(resume_text)
            st.header("Strengths and Weaknesses Analysis")
            st.write(strengths_weaknesses)


if __name__== "__main__":
    main()